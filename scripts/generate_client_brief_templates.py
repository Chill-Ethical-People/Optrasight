from __future__ import annotations

from email.message import EmailMessage
from email.policy import default
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "deliverables"
DOCX_PATH = OUTPUT / "OptraSight_Client_Threat_Intelligence_Brief_Template.docx"
EML_PATH = OUTPUT / "OptraSight_Client_Threat_Intelligence_Brief_Template.eml"
LOGO_PATH = ROOT / "client/public/brand/optrasight-smooth-mark-light-256.png"

BRAND = "4F46E5"
SIGNAL = "0891B2"
INK = "111827"
MUTED = "667085"
SOFT = "EEF0FE"
RED = "B42318"
RED_SOFT = "FEE4E2"
AMBER = "B54708"
AMBER_SOFT = "FEF0C7"
BLUE = "175CD3"
BLUE_SOFT = "D1E9FF"
GRAY_SOFT = "F2F4F7"


def shade(cell, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    fill = tc_pr.find(qn("w:shd"))
    if fill is None:
        fill = OxmlElement("w:shd")
        tc_pr.append(fill)
    fill.set(qn("w:fill"), color)


def set_cell_margins(cell, top=120, start=150, bottom=120, end=150) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run(run, size=10, bold=False, color=INK, font="Aptos") -> None:
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def add_text(paragraph, text: str, **kwargs):
    run = paragraph.add_run(text)
    set_run(run, **kwargs)
    return run


def add_label_value(cell, label: str, value: str, accent: str = INK) -> None:
    set_cell_margins(cell, top=130, bottom=130)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    add_text(p, label.upper(), size=7.5, bold=True, color=MUTED)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    add_text(p2, value, size=12, bold=True, color=accent)


def add_section_bar(doc: Document, title: str, subtitle: str, color: str, soft_color: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    cell = table.cell(0, 0)
    shade(cell, soft_color)
    set_cell_margins(cell, top=140, bottom=140, start=180, end=180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    add_text(p, title, size=12, bold=True, color=color)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    add_text(p2, subtitle, size=8.5, color=MUTED)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_finding(doc: Document, number: int, title: str, risk: str, status: str, color: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    add_text(p, f"{number}. {title}", size=11, bold=True)

    meta = doc.add_table(rows=1, cols=4)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta.autofit = True
    labels = [
        ("Risk", risk),
        ("Intel status", status),
        ("Published", "[DD MMM YYYY]"),
        ("Affected scope", "[Technology / function]"),
    ]
    for cell, (label, value) in zip(meta.rows[0].cells, labels):
        shade(cell, GRAY_SOFT)
        set_cell_margins(cell, top=90, bottom=90)
        p1 = cell.paragraphs[0]
        add_text(p1, label.upper(), size=6.8, bold=True, color=MUTED)
        p2 = cell.add_paragraph()
        add_text(p2, value, size=8.5, bold=label == "Risk", color=color if label == "Risk" else INK)

    details = [
        ("Why this matters", "[Explain the client-specific exposure or business impact in two concise sentences.]"),
        ("Recommended action", "[Provide one concrete validation, mitigation, detection, or monitoring action.]"),
        ("Owner and timing", "[SOC / VM / IR / Detection Engineering] - [Immediate / 24 hours / this week / monitor]"),
        ("Source", "[Publisher - advisory title and URL]"),
    ]
    for label, value in details:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(1)
        add_text(p, f"{label}: ", size=9, bold=True, color=color)
        add_text(p, value, size=9)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def build_docx() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(9.5)
    styles["Normal"].font.color.rgb = RGBColor.from_string(INK)

    header = section.header
    header_table = header.add_table(rows=1, cols=2, width=Inches(7.0))
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_table.columns[0].width = Inches(4.9)
    header_table.columns[1].width = Inches(2.1)
    left = header_table.cell(0, 0)
    right = header_table.cell(0, 1)
    if LOGO_PATH.exists():
        left.paragraphs[0].add_run().add_picture(str(LOGO_PATH), width=Inches(0.36))
    add_text(left.paragraphs[0], "  OptraSight", size=13, bold=True, color=INK)
    add_text(left.paragraphs[0], "  THREAT INTELLIGENCE", size=7.5, bold=True, color=BRAND)
    right.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_text(right.paragraphs[0], "DRAFT - CLIENT REVIEW", size=8, bold=True, color=RED)

    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(footer_p, "CONFIDENTIAL  |  OptraSight Threat Intelligence  |  Analyst approval required before distribution", size=7.5, bold=True, color=MUTED)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    add_text(p, "CLIENT RISK BRIEF", size=20, bold=True, color=INK)
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(10)
    add_text(p2, "Action-oriented risk summary for security leadership and operations", size=9.5, color=MUTED)

    facts = doc.add_table(rows=2, cols=4)
    facts.alignment = WD_TABLE_ALIGNMENT.CENTER
    values = [
        ("Client", "[Client name]"),
        ("Reporting period", "[Start] - [End]"),
        ("Cadence", "[Weekly / Monthly]"),
        ("Overall risk", "HIGH"),
        ("Prepared by", "OptraSight TI"),
        ("Issue date", "[DD MMM YYYY]"),
        ("Risk trend", "[Increasing / Stable]"),
        ("Reference", "[BRIEF-YYYY-NNN]"),
    ]
    for cell, (label, value) in zip([c for row in facts.rows for c in row.cells], values):
        shade(cell, SOFT if label in {"Client", "Overall risk"} else "F8FAFC")
        add_label_value(cell, label, value, RED if label == "Overall risk" else BRAND if label == "Client" else INK)

    doc.add_paragraph()
    subject_table = doc.add_table(rows=1, cols=1)
    subject_cell = subject_table.cell(0, 0)
    shade(subject_cell, "F8FAFC")
    set_cell_margins(subject_cell, top=120, bottom=120, start=180, end=180)
    subject_p = subject_cell.paragraphs[0]
    add_text(subject_p, "EMAIL SUBJECT  ", size=7.5, bold=True, color=MUTED)
    add_text(subject_p, "[OptraSight Threat Intelligence] [Client] | [Weekly / Monthly] Risk Brief | [DD MMM YYYY]", size=9, bold=True)
    greeting = doc.add_paragraph()
    greeting.paragraph_format.space_before = Pt(10)
    greeting.paragraph_format.space_after = Pt(8)
    add_text(greeting, "Hello [Client] Security Team,", size=9.5)

    add_section_bar(doc, "EXECUTIVE SUMMARY", "Aggregate risk posture and the decisions that need attention", BRAND, SOFT)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    add_text(p, "During this reporting period, OptraSight reviewed ", size=9.5)
    add_text(p, "[total] client-relevant intelligence items", size=9.5, bold=True)
    add_text(p, ". The primary concern is [one-sentence description of the most important change in risk]. We recommend [highest-priority client action].", size=9.5)

    metrics = doc.add_table(rows=1, cols=4)
    metrics.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell, data in zip(metrics.rows[0].cells, [
        ("TIER 1", "[0]", "Action required", RED, RED_SOFT),
        ("TIER 2", "[0]", "Priority review", AMBER, AMBER_SOFT),
        ("TIER 3", "[0]", "Monitor and plan", BLUE, BLUE_SOFT),
        ("FYI", "[0]", "Awareness only", MUTED, GRAY_SOFT),
    ]):
        label, count, note, color, soft = data
        shade(cell, soft)
        set_cell_margins(cell, top=110, bottom=110)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_text(p, label, size=7.5, bold=True, color=color)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_text(p2, count, size=16, bold=True, color=color)
        p3 = cell.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_text(p3, note, size=7.5, color=MUTED)

    add_section_bar(doc, "TIER 1 - ACTION REQUIRED", "Confirmed or imminent client risk requiring immediate operational action", RED, RED_SOFT)
    add_finding(doc, 1, "[Critical threat or advisory title]", "CRITICAL", "Escalated", RED)

    add_section_bar(doc, "TIER 2 - PRIORITY REVIEW", "Credible client exposure requiring validation or remediation this week", AMBER, AMBER_SOFT)
    add_finding(doc, 1, "[High-priority threat or advisory title]", "HIGH", "Assessed", AMBER)

    add_section_bar(doc, "TIER 3 - MONITOR AND PLAN", "Relevant development for monitoring, backlog planning, or control improvement", BLUE, BLUE_SOFT)
    add_finding(doc, 1, "[Medium-priority threat or advisory title]", "MEDIUM", "Triaged", BLUE)

    add_section_bar(doc, "FYI - SITUATIONAL AWARENESS", "Contextual intelligence with no immediate client action", MUTED, GRAY_SOFT)
    for text in (
        "[Finding title] - [One-sentence summary and client relevance]. [Source URL]",
        "[Finding title] - [One-sentence summary and client relevance]. [Source URL]",
    ):
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        add_text(p, text, size=9)

    add_section_bar(doc, "RECOMMENDED ACTION SUMMARY", "Prioritised handoff for client security teams", SIGNAL, "CFFAFE")
    actions = doc.add_table(rows=1, cols=4)
    actions.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Priority", "Action", "Suggested owner", "Target timing"]
    for cell, text in zip(actions.rows[0].cells, headers):
        shade(cell, INK)
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        add_text(p, text, size=8, bold=True, color="FFFFFF")
    set_repeat_table_header(actions.rows[0])
    for values in (
        ("1", "[Immediate Tier 1 action]", "[IR / SOC / VM]", "Immediate / 24 hours"),
        ("2", "[Tier 2 validation or detection action]", "[Detection Engineering]", "This week"),
        ("3", "[Tier 3 monitoring action]", "[Threat Intelligence]", "Next review"),
    ):
        cells = actions.add_row().cells
        for cell, value in zip(cells, values):
            set_cell_margins(cell)
            add_text(cell.paragraphs[0], value, size=8.5, bold=value == values[0])

    doc.add_paragraph()
    p = doc.add_paragraph()
    add_text(p, "Supporting material: ", size=9, bold=True)
    add_text(p, "[IOC count] indicators | [CVE list] | Detection guidance [included / available on request]", size=9)
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(8)
    add_text(p2, "Please contact the threat-intelligence team for supporting indicators, detection queries, or additional analysis.", size=9)
    p3 = doc.add_paragraph()
    add_text(p3, "Regards,\nOptraSight Threat Intelligence\n[Contact details]", size=9, bold=True, color=BRAND)

    OUTPUT.mkdir(parents=True, exist_ok=True)
    doc.save(DOCX_PATH)


def build_eml() -> None:
    subject = "[DRAFT] [OptraSight Threat Intelligence] [Client] | Weekly Risk Brief | [DD MMM YYYY]"
    plain = """Hello [Client] Security Team,

EXECUTIVE SUMMARY
Overall risk: HIGH | Trend: [Increasing / Stable / Decreasing]
Tier 1 Action Required: [0] | Tier 2 Priority Review: [0] | Tier 3 Monitor and Plan: [0] | FYI: [0]

During this reporting period, OptraSight reviewed [total] client-relevant intelligence items. The primary concern is [summary]. We recommend [highest-priority action].

TIER 1 - ACTION REQUIRED
[Finding title]
Risk: Critical | Status: Escalated | Timing: Immediate / 24 hours
Why this matters: [Client-specific exposure or business impact.]
Required action: [Concrete mitigation or validation step.]
Owner: [IR / SOC / Vulnerability Management]
Source: [Publisher and URL]

TIER 2 - PRIORITY REVIEW
[Finding title]
Risk: High | Status: Assessed | Timing: This week
Client relevance: [Why this may affect the client.]
Recommended action: [Validation, patch review, detection, or hunt.]
Source: [Publisher and URL]

TIER 3 - MONITOR AND PLAN
[Finding title] - [Relevance and monitoring action.]

FYI - SITUATIONAL AWARENESS
- [Finding title] - [One-sentence summary and source URL.]

RECOMMENDED ACTION SUMMARY
1. [Immediate Tier 1 action] - [Owner] - [Timing]
2. [Tier 2 validation action] - [Owner] - [Timing]
3. [Tier 3 monitoring action] - [Owner] - [Timing]

Supporting indicators: [count]
Relevant CVEs: [CVE list]
Detection guidance: [Included / Available on request]

Regards,
OptraSight Threat Intelligence
[Contact details]

DRAFT - Analyst approval required before distribution.
"""
    html = """<!doctype html>
<html><body style="margin:0;background:#f2f4f7;font-family:Arial,Helvetica,sans-serif;color:#111827">
<table role="presentation" width="100%" cellspacing="0" cellpadding="0" style="background:#f2f4f7"><tr><td align="center" style="padding:28px 12px">
<table role="presentation" width="700" cellspacing="0" cellpadding="0" style="width:100%;max-width:700px;background:#ffffff;border:1px solid #d0d5dd">
<tr><td style="padding:22px 28px;border-top:5px solid #4F46E5;border-bottom:1px solid #e4e7ec">
  <table role="presentation" width="100%"><tr><td><div style="font-size:22px;font-weight:700">Optra<span style="color:#4F46E5">Sight</span></div><div style="font-size:11px;color:#667085;letter-spacing:1px">THREAT INTELLIGENCE</div></td>
  <td align="right"><span style="display:inline-block;padding:6px 10px;background:#FEE4E2;color:#B42318;font-size:11px;font-weight:700">DRAFT - CLIENT REVIEW</span></td></tr></table>
</td></tr>
<tr><td style="padding:28px">
  <div style="font-size:12px;color:#4F46E5;font-weight:700;letter-spacing:1px">WEEKLY CLIENT RISK BRIEF</div>
  <h1 style="margin:6px 0 4px;font-size:26px;line-height:1.25">[Client name]</h1>
  <div style="color:#667085;font-size:13px">Reporting period: [Start] - [End] &nbsp;|&nbsp; Issued: [DD MMM YYYY]</div>

  <table role="presentation" width="100%" cellspacing="0" cellpadding="0" style="margin-top:22px"><tr>
    <td width="25%" style="padding:12px;background:#FEE4E2;border-right:6px solid #fff"><div style="font-size:10px;color:#B42318;font-weight:700">TIER 1</div><div style="font-size:22px;font-weight:700;color:#B42318">[0]</div><div style="font-size:11px;color:#667085">Action required</div></td>
    <td width="25%" style="padding:12px;background:#FEF0C7;border-right:6px solid #fff"><div style="font-size:10px;color:#B54708;font-weight:700">TIER 2</div><div style="font-size:22px;font-weight:700;color:#B54708">[0]</div><div style="font-size:11px;color:#667085">Priority review</div></td>
    <td width="25%" style="padding:12px;background:#D1E9FF;border-right:6px solid #fff"><div style="font-size:10px;color:#175CD3;font-weight:700">TIER 3</div><div style="font-size:22px;font-weight:700;color:#175CD3">[0]</div><div style="font-size:11px;color:#667085">Monitor and plan</div></td>
    <td width="25%" style="padding:12px;background:#F2F4F7"><div style="font-size:10px;color:#667085;font-weight:700">FYI</div><div style="font-size:22px;font-weight:700">[0]</div><div style="font-size:11px;color:#667085">Awareness only</div></td>
  </tr></table>

  <div style="margin-top:22px;padding:18px;background:#EEF0FE;border-left:4px solid #4F46E5">
    <div style="font-size:11px;color:#4F46E5;font-weight:700">EXECUTIVE SUMMARY</div>
    <p style="margin:8px 0 0;font-size:14px;line-height:1.6">Overall risk is <b>HIGH</b> with a <b>[stable / increasing]</b> trend. During this period, OptraSight reviewed <b>[total]</b> client-relevant items. The primary concern is [one-sentence risk summary]. We recommend [highest-priority action].</p>
  </div>

  <div style="margin-top:28px;border-top:4px solid #B42318;background:#fff7f6;padding:18px">
    <div style="font-size:15px;color:#B42318;font-weight:700">TIER 1 - ACTION REQUIRED</div>
    <div style="margin-top:4px;font-size:11px;color:#667085">Confirmed or imminent risk requiring immediate operational action</div>
    <h2 style="margin:18px 0 8px;font-size:18px">[Critical threat or advisory title]</h2>
    <div style="font-size:12px"><b style="color:#B42318">CRITICAL</b> &nbsp;|&nbsp; Escalated &nbsp;|&nbsp; [Affected technology]</div>
    <p style="font-size:14px;line-height:1.55"><b>Why this matters:</b> [Client-specific exposure or business impact.]</p>
    <p style="font-size:14px;line-height:1.55"><b>Required action:</b> [Concrete mitigation or validation step.]</p>
    <p style="font-size:14px;line-height:1.55"><b>Owner and timing:</b> [IR / SOC / VM] - Immediate / within 24 hours</p>
    <p style="font-size:13px"><a href="https://example.com/advisory" style="color:#4F46E5">[Publisher - advisory title]</a></p>
  </div>

  <div style="margin-top:18px;border-top:4px solid #B54708;background:#fffbeb;padding:18px">
    <div style="font-size:15px;color:#B54708;font-weight:700">TIER 2 - PRIORITY REVIEW</div>
    <h2 style="margin:14px 0 8px;font-size:18px">[High-priority threat or advisory title]</h2>
    <p style="font-size:14px;line-height:1.55"><b>Client relevance:</b> [Why this may affect the client.]</p>
    <p style="font-size:14px;line-height:1.55"><b>Recommended action:</b> [Validation, patch review, detection, or hunt.]</p>
    <p style="font-size:14px;line-height:1.55"><b>Owner and timing:</b> [Security function] - This week</p>
  </div>

  <div style="margin-top:18px;border-top:4px solid #175CD3;background:#f5faff;padding:18px">
    <div style="font-size:15px;color:#175CD3;font-weight:700">TIER 3 - MONITOR AND PLAN</div>
    <p style="font-size:14px;line-height:1.55"><b>[Finding title]</b> - [Client relevance and monitoring or planning action.]</p>
  </div>

  <div style="margin-top:18px;background:#F2F4F7;padding:18px">
    <div style="font-size:15px;font-weight:700">FYI - SITUATIONAL AWARENESS</div>
    <ul style="font-size:14px;line-height:1.65;padding-left:20px"><li><b>[Finding title]</b> - [One-sentence summary and relevance].</li><li><b>[Finding title]</b> - [One-sentence summary and relevance].</li></ul>
  </div>

  <h2 style="margin:28px 0 10px;font-size:17px">Recommended action summary</h2>
  <table role="presentation" width="100%" cellspacing="0" cellpadding="8" style="font-size:13px;border-collapse:collapse">
    <tr style="background:#111827;color:#fff"><th align="left">Priority</th><th align="left">Action</th><th align="left">Owner</th><th align="left">Timing</th></tr>
    <tr><td style="border-bottom:1px solid #e4e7ec">1</td><td style="border-bottom:1px solid #e4e7ec">[Immediate Tier 1 action]</td><td style="border-bottom:1px solid #e4e7ec">[IR / SOC / VM]</td><td style="border-bottom:1px solid #e4e7ec">24 hours</td></tr>
    <tr><td style="border-bottom:1px solid #e4e7ec">2</td><td style="border-bottom:1px solid #e4e7ec">[Tier 2 validation action]</td><td style="border-bottom:1px solid #e4e7ec">[Detection Engineering]</td><td style="border-bottom:1px solid #e4e7ec">This week</td></tr>
    <tr><td>3</td><td>[Tier 3 monitoring action]</td><td>[Threat Intelligence]</td><td>Next review</td></tr>
  </table>

  <p style="margin-top:24px;font-size:13px;line-height:1.6">Supporting indicators: <b>[count]</b> &nbsp;|&nbsp; Relevant CVEs: <b>[CVE list]</b> &nbsp;|&nbsp; Detection guidance: <b>[included / available on request]</b></p>
  <p style="font-size:14px;line-height:1.6">Please contact the threat-intelligence team for supporting indicators, detection queries, or additional analysis.</p>
  <p style="font-size:14px;line-height:1.6"><b>Regards,<br>OptraSight Threat Intelligence</b><br>[Contact details]</p>
</td></tr>
<tr><td style="padding:18px 28px;background:#111827;color:#98A2B3;font-size:11px;line-height:1.5">CONFIDENTIAL - This draft requires analyst approval before external distribution. Do not forward without authorisation.</td></tr>
</table></td></tr></table></body></html>"""

    msg = EmailMessage(policy=default)
    msg["Subject"] = subject
    msg["From"] = "OptraSight Threat Intelligence <threat-intel@example.com>"
    msg["To"] = "[Client security distribution list]"
    msg["X-Unsent"] = "1"
    msg.set_content(plain)
    msg.add_alternative(html, subtype="html")
    OUTPUT.mkdir(parents=True, exist_ok=True)
    EML_PATH.write_bytes(msg.as_bytes())


if __name__ == "__main__":
    build_docx()
    build_eml()
    print(DOCX_PATH)
    print(EML_PATH)
